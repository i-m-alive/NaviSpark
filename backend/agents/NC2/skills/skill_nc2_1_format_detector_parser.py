"""
Skill NC2.1 — Format Detector & Parser

Detects the file format of the uploaded checklist and extracts its raw content
into a normalised intermediate structure. This is the gateway skill — all other
NC2 skills receive its output.

Supported formats: Excel (.xlsx, .xlsm), CSV (.csv), DOCX (.docx), PDF (.pdf).
"""

from __future__ import annotations

import csv
import logging
import mimetypes
import os
import re
from typing import Any

logger = logging.getLogger(__name__)

_BULLET_CHARS = frozenset({"•", "-", "*", "–", "—", "·", "▪", "◦"})
_NUMBERED_LIST_RE = re.compile(r"^\d+[\.\)]\s")


class FormatDetectorParser:
    """Detects the file format of a checklist and extracts its raw row content.

    Each raw_row in the returned list contains at minimum:
      - row_index (int): position in the source
      - raw_text (str): all cell values joined by " | "
      - source (str): human-readable location descriptor
      - fields (dict): original key-value pairs where available

    The returned dict also carries a top-level ``headers`` list (the first-row
    column names) so downstream skills can perform field mapping.
    """

    def detect_format(self, file_path: str) -> str:
        """Detect the checklist file format from its extension or MIME type.

        Args:
            file_path: Absolute path to the checklist file.

        Returns:
            One of: "xlsx", "csv", "docx", "pdf".

        Raises:
            ValueError: If the format cannot be determined or is unsupported.
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext in (".xlsx", ".xlsm"):
            return "xlsx"
        if ext == ".csv":
            return "csv"
        if ext == ".docx":
            return "docx"
        if ext == ".pdf":
            return "pdf"

        mime, _ = mimetypes.guess_type(file_path)
        if mime:
            if "spreadsheet" in mime or "excel" in mime:
                return "xlsx"
            if mime == "text/csv":
                return "csv"
            if "wordprocessingml" in mime or "msword" in mime:
                return "docx"
            if mime == "application/pdf":
                return "pdf"

        raise ValueError(
            f"Unsupported checklist file format: {file_path}. "
            "Supported: xlsx, csv, docx, pdf"
        )

    def parse(self, file_path: str) -> dict[str, Any]:
        """Detect format and dispatch to the appropriate internal parser.

        Args:
            file_path: Absolute path to the checklist file.

        Returns:
            A dict with keys:
              - format (str): detected format string.
              - raw_rows (list[dict]): one dict per data row.
              - sheet_names (list[str]): sheet names for xlsx; empty otherwise.
              - parse_warnings (list[str]): non-fatal issues encountered.
              - headers (list[str]): column headers from the source file.
        """
        fmt = self.detect_format(file_path)
        logger.debug("NC2.1 detected format=%s for %s", fmt, os.path.basename(file_path))

        if fmt == "xlsx":
            return self._parse_xlsx(file_path)
        if fmt == "csv":
            return self._parse_csv(file_path)
        if fmt == "docx":
            return self._parse_docx(file_path)
        return self._parse_pdf(file_path)

    def _parse_xlsx(self, file_path: str) -> dict[str, Any]:
        """Parse an Excel workbook into raw rows.

        All sheets are iterated. The first non-empty row with 2+ populated cells
        is treated as the header row for that sheet. Merged cells are forward-filled.

        Args:
            file_path: Path to the .xlsx or .xlsm file.

        Returns:
            Parsed dict matching the parse() return schema.
        """
        try:
            import openpyxl
        except ImportError as exc:
            raise RuntimeError("openpyxl is required to parse Excel files.") from exc

        raw_rows: list[dict] = []
        parse_warnings: list[str] = []
        sheet_names: list[str] = []
        first_sheet_headers: list[str] = []

        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            sheet_names = wb.sheetnames

            for sheet_name in sheet_names:
                ws = wb[sheet_name]
                sheet_headers: list[str] = []
                header_found = False
                last_values: dict[int, Any] = {}
                data_count = 0

                for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                    cells = [c if c is not None else "" for c in row]

                    for col_idx, val in enumerate(cells):
                        if val != "":
                            last_values[col_idx] = val
                        elif col_idx in last_values:
                            cells[col_idx] = last_values[col_idx]

                    non_empty = [str(c).strip() for c in cells if str(c).strip()]
                    if not non_empty:
                        continue

                    if not header_found and len(non_empty) >= 2:
                        sheet_headers = [str(c).strip() for c in cells]
                        if not first_sheet_headers:
                            first_sheet_headers = [h for h in sheet_headers if h]
                        header_found = True
                        continue

                    raw_text = " | ".join(str(c).strip() for c in cells if str(c).strip())
                    fields: dict[str, str] = {}
                    if sheet_headers:
                        for h, c in zip(sheet_headers, cells):
                            if h:
                                fields[h] = str(c).strip() if c else ""

                    raw_rows.append({
                        "row_index": row_idx,
                        "raw_text": raw_text,
                        "source": f"{sheet_name}:Row{row_idx}",
                        "fields": fields,
                    })
                    data_count += 1

                if data_count == 0:
                    parse_warnings.append(
                        f"Sheet '{sheet_name}' has no data rows — skipped"
                    )

            wb.close()
        except Exception as exc:
            logger.error("NC2.1 xlsx parsing error: %s", exc)
            raise RuntimeError(f"Failed to parse Excel file: {exc}") from exc

        return {
            "format": "xlsx",
            "raw_rows": raw_rows,
            "sheet_names": sheet_names,
            "parse_warnings": parse_warnings,
            "headers": first_sheet_headers,
        }

    def _parse_csv(self, file_path: str) -> dict[str, Any]:
        """Parse a CSV file into raw rows using auto-detected delimiter.

        Args:
            file_path: Path to the .csv file.

        Returns:
            Parsed dict matching the parse() return schema.
        """
        raw_rows: list[dict] = []
        parse_warnings: list[str] = []

        try:
            with open(file_path, encoding="utf-8-sig", errors="replace") as f:
                first_line = f.readline()
        except Exception as exc:
            raise RuntimeError(f"Cannot read CSV file: {exc}") from exc

        if "\t" in first_line and "," not in first_line:
            delimiter = "\t"
        elif ";" in first_line and "," not in first_line:
            delimiter = ";"
        else:
            delimiter = ","

        try:
            with open(
                file_path, encoding="utf-8-sig", errors="replace", newline=""
            ) as f:
                reader = csv.DictReader(f, delimiter=delimiter)
                headers: list[str] = list(reader.fieldnames or [])

                for row_idx, row in enumerate(reader, start=2):
                    values = [v.strip() if v else "" for v in row.values()]
                    if all(v == "" for v in values):
                        continue

                    raw_text = " | ".join(v for v in values if v)
                    raw_rows.append({
                        "row_index": row_idx,
                        "raw_text": raw_text,
                        "source": f"CSV:Row{row_idx}",
                        "fields": {k: (v.strip() if v else "") for k, v in row.items()},
                    })
        except Exception as exc:
            raise RuntimeError(f"Failed to parse CSV file: {exc}") from exc

        if not raw_rows:
            parse_warnings.append("CSV file has a header but no data rows")

        logger.debug("NC2.1 CSV parsed %d rows, headers=%s", len(raw_rows), headers)

        return {
            "format": "csv",
            "raw_rows": raw_rows,
            "sheet_names": [],
            "parse_warnings": parse_warnings,
            "headers": headers,
        }

    def _parse_docx(self, file_path: str) -> dict[str, Any]:
        """Parse a DOCX file extracting tables and bulleted/numbered paragraphs.

        Tables: each data row (after the header row) becomes a raw_row.
        Bulleted/numbered paragraphs: each list item becomes a raw_row.

        Args:
            file_path: Path to the .docx file.

        Returns:
            Parsed dict matching the parse() return schema.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to parse DOCX files.") from exc

        raw_rows: list[dict] = []
        parse_warnings: list[str] = []
        first_table_headers: list[str] = []

        try:
            doc = Document(file_path)

            for table_idx, table in enumerate(doc.tables):
                if not table.rows:
                    continue
                header_cells = [c.text.strip() for c in table.rows[0].cells]
                if not first_table_headers and any(header_cells):
                    first_table_headers = [h for h in header_cells if h]

                for row_idx, row in enumerate(table.rows[1:], start=1):
                    cell_texts = [c.text.strip() for c in row.cells]
                    if not any(cell_texts):
                        continue
                    raw_text = " | ".join(t for t in cell_texts if t)
                    fields: dict[str, str] = {}
                    if header_cells:
                        for h, t in zip(header_cells, cell_texts):
                            if h:
                                fields[h] = t
                    raw_rows.append({
                        "row_index": row_idx,
                        "raw_text": raw_text,
                        "source": f"DOCX:Table{table_idx}:Row{row_idx}",
                        "fields": fields,
                    })

            list_count = 0
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                is_list_style = "List" in (para.style.name or "")
                starts_with_bullet = text[0] in _BULLET_CHARS
                is_numbered = bool(_NUMBERED_LIST_RE.match(text))

                if is_list_style or starts_with_bullet or is_numbered:
                    clean = text.lstrip("".join(_BULLET_CHARS)).strip()
                    clean = _NUMBERED_LIST_RE.sub("", clean).strip()
                    raw_rows.append({
                        "row_index": para_idx,
                        "raw_text": clean,
                        "source": f"DOCX:Bullet:Para{para_idx}",
                        "fields": {},
                    })
                    list_count += 1

            if not doc.tables and list_count == 0:
                parse_warnings.append(
                    "DOCX has no tables or list paragraphs — attempting full paragraph extraction"
                )
                for para_idx, para in enumerate(doc.paragraphs):
                    text = para.text.strip()
                    if text:
                        raw_rows.append({
                            "row_index": para_idx,
                            "raw_text": text,
                            "source": f"DOCX:Para{para_idx}",
                            "fields": {},
                        })

        except Exception as exc:
            logger.error("NC2.1 DOCX parsing error: %s", exc)
            raise RuntimeError(f"Failed to parse DOCX file: {exc}") from exc

        logger.debug("NC2.1 DOCX parsed %d rows", len(raw_rows))

        return {
            "format": "docx",
            "raw_rows": raw_rows,
            "sheet_names": [],
            "parse_warnings": parse_warnings,
            "headers": first_table_headers,
        }

    def _parse_pdf(self, file_path: str) -> dict[str, Any]:
        """Parse a PDF file using pypdf for text extraction.

        If pdfplumber is available, also attempts table extraction for richer output.

        Args:
            file_path: Path to the .pdf file.

        Returns:
            Parsed dict matching the parse() return schema.
        """
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("pypdf is required to parse PDF files.") from exc

        try:
            import pdfplumber as _pdfplumber  # type: ignore
        except ImportError:
            _pdfplumber = None  # type: ignore

        raw_rows: list[dict] = []
        parse_warnings: list[str] = []
        total_text_len = 0

        try:
            reader = PdfReader(file_path)

            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    page_text = page.extract_text() or ""
                except Exception as exc:
                    logger.warning("NC2.1 PDF page %d text extraction failed: %s", page_num, exc)
                    page_text = ""

                total_text_len += len(page_text)
                lines = page_text.splitlines()

                for line_idx, line in enumerate(lines, start=1):
                    stripped = line.strip()
                    if not stripped or len(stripped) < 5:
                        continue
                    raw_rows.append({
                        "row_index": line_idx,
                        "raw_text": stripped,
                        "source": f"PDF:Page{page_num}:Line{line_idx}",
                        "fields": {},
                    })

            if total_text_len < 200:
                parse_warnings.append(
                    "PDF text extraction returned very little content "
                    "— PDF may be scanned or image-based"
                )

        except Exception as exc:
            logger.error("NC2.1 PDF pypdf parsing error: %s", exc)
            raise RuntimeError(f"Failed to parse PDF file: {exc}") from exc

        if _pdfplumber is not None:
            table_count = 0
            try:
                with _pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, start=1):
                        try:
                            tables = page.extract_tables() or []
                        except Exception:
                            tables = []
                        for t_idx, table in enumerate(tables):
                            for r_idx, row in enumerate(table):
                                if not row or not any(c for c in row if c):
                                    continue
                                raw_text = " | ".join(str(c).strip() for c in row if c)
                                raw_rows.append({
                                    "row_index": r_idx,
                                    "raw_text": raw_text,
                                    "source": f"PDF:pdfplumber:Page{page_num}:Table{t_idx}:Row{r_idx}",
                                    "fields": {},
                                })
                                table_count += 1
                if table_count:
                    parse_warnings.append(
                        f"pdfplumber table extraction used for {table_count} tables"
                    )
            except Exception as exc:
                logger.warning("NC2.1 pdfplumber extraction failed: %s", exc)

        logger.debug("NC2.1 PDF parsed %d rows", len(raw_rows))

        return {
            "format": "pdf",
            "raw_rows": raw_rows,
            "sheet_names": [],
            "parse_warnings": parse_warnings,
            "headers": [],
        }
